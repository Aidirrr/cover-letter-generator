import streamlit as st
import google.generativeai as genai
import PyPDF2
from fpdf import FPDF
from docx import Document
from datetime import datetime

# Streamlit App Title
st.title("📝 AI-Powered Cover Letter Generator")

# Sidebar for API Key
st.sidebar.header("🔑 Settings")
api_key = st.sidebar.text_input("Enter Gemini API Key", type="password")


# Function to extract text from PDF
def extract_text_from_pdf(uploaded_file):
    pdf_reader = PyPDF2.PdfReader(uploaded_file)
    text = "".join([page.extract_text() for page in pdf_reader.pages if page.extract_text()])
    return text


# Function to create PDF
def create_pdf(text):
    pdf = FPDF()
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.add_page()
    pdf.set_font("Arial", size=12)
    pdf.multi_cell(0, 10, text)
    return pdf


# Function to create DOCX
def create_docx(full_name, address, email, phone, linkedin, github, today_date, company, location, position,
                cover_letter):
    doc = Document()
    doc.add_paragraph(full_name)
    doc.add_paragraph(address)
    doc.add_paragraph(f"{email} | {phone}")
    doc.add_paragraph(today_date)
    doc.add_paragraph("Hiring Manager")
    doc.add_paragraph(company)
    doc.add_paragraph(location)
    doc.add_paragraph(f"Subject: Application for {position}")
    doc.add_paragraph(cover_letter)
    return doc


# User Inputs
st.header("👤 Personal Information")
col1, col2 = st.columns(2)

with col1:
    full_name = st.text_input("Full Name")
    address = st.text_area("Address")
    email = st.text_input("Email")

with col2:
    phone = st.text_input("Phone Number")
    linkedin = st.text_input("LinkedIn Profile URL (Optional)")
    github = st.text_input("GitHub Profile URL (Optional)")

today_date = datetime.today().strftime('%d %B %Y')

st.header("💼 Job Details")
col3, col4 = st.columns(2)

with col3:
    company = st.text_input("Company Name")
    position = st.text_input("Position")

with col4:
    location = st.text_input("Location")
    job_description = st.text_area("Job Description")

# Additional Information
st.header("📝 Additional Information")
additional_info = st.text_area(
    "Provide additional details (e.g., applying for an internship, internship duration, etc.)")

st.header("📄 Upload Your Resume")
uploaded_resume = st.file_uploader("Upload your resume (PDF only)", type=["pdf"])
resume_text = ""
if uploaded_resume:
    resume_text = extract_text_from_pdf(uploaded_resume)
    st.success("✅ Resume uploaded successfully!")

# Generate Cover Letter
if st.button("✨ Generate Cover Letter"):
    if not api_key:
        st.error("🚨 Please enter your Gemini API Key in the sidebar.")
    elif not full_name or not email or not phone or not company or not position or not job_description:
        st.error("🚨 Please fill in all required fields.")
    else:
        with st.spinner("🔍 Generating your cover letter..."):
            genai.configure(api_key=api_key)
            model = genai.GenerativeModel("gemini-pro")

            prompt = f"""
                    Write a professional and tailored cover letter using ONLY the information provided below. Do not add any details, experiences, or company names that are not explicitly mentioned in the input. Focus on highlighting the candidate's skills, experiences, and motivations based on the provided resume content and job description.

                    ---

                    **Personal Details:**
                    - Name: {full_name}
                    - Address: {address}
                    - Email: {email}
                    - Phone: {phone}
                    - LinkedIn: {linkedin if linkedin else 'N/A'}
                    - GitHub: {github if github else 'N/A'}
                    - Date: {today_date}

                    **Job Details:**
                    - Position: {position}
                    - Company: {company}
                    - Location: {location}
                    - Job Description: {job_description}

                    **Additional Information:**
                    {additional_info if additional_info else 'No additional information provided.'}

                    **Resume Content:**
                    {resume_text}

                    ---

                    **Format the cover letter as follows:**

                    {full_name}
                    {address}
                    {email} | {phone}

                    {today_date}

                    Hiring Manager
                    {company}
                    {location}

                    Subject: Application for {position}

                    [Cover Letter Content]

                    - Start with a strong opening paragraph that mentions the position you are applying for and a brief statement about why you are excited about the opportunity.
                    - In the second paragraph, highlight your relevant skills, experiences, and achievements from the resume content that align with the job description. Do not invent any experiences or company names.
                    - In the third paragraph, mention any additional information provided (e.g., internship duration, specific motivations) and explain how it makes you a good fit for the role.
                    - Conclude with a professional closing statement expressing your enthusiasm for the opportunity and your availability for an interview.

                    Best regards,
                    {full_name}
            """

            response = model.generate_content(prompt)
            cover_letter = response.text

            st.subheader("📜 Generated Cover Letter")
            st.text_area("", cover_letter, height=300)

            # Create PDF
            pdf = create_pdf(f"{full_name}\n{address}\n{email} | {phone}\n{today_date}\n\n{cover_letter}")
            pdf_bytes = pdf.output(dest="S").encode("latin1")
            st.download_button("📥 Download as PDF", pdf_bytes, file_name="cover_letter.pdf", mime="application/pdf")

            # Create DOCX
            docx = create_docx(full_name, address, email, phone, linkedin, github, today_date, company, location,
                               position, cover_letter)
            docx.save("cover_letter.docx")
            with open("cover_letter.docx", "rb") as f:
                st.download_button("📥 Download as DOCX", f, file_name="cover_letter.docx",
                                   mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")

# Footer
st.markdown("---")
st.markdown("Made with ❤️ by Khaidhir")